 
import os
from PyPDF2 import PdfReader
from docx import Document
import re

def extract_text_from_resume(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text

    else:
        raise ValueError("Unsupported file format")

def get_recommendations(resume_text: str) -> list:
    resume_text = resume_text.lower()

    career_paths = {
        "data science": ["python", "machine learning", "data analysis", "pandas", "numpy", "scikit-learn", "tensorflow"],
        "web development": ["html", "css", "javascript", "react", "angular", "node", "frontend", "backend"],
        "android development": ["kotlin", "java", "android studio", "xml"],
        "cloud engineer": ["aws", "azure", "gcp", "cloud", "docker", "kubernetes", "devops"],
        "ai engineer": ["deep learning", "nlp", "transformers", "pytorch", "tensorflow", "huggingface", "llm"]
    }

    recommendations = []

    for career, keywords in career_paths.items():
        if any(skill in resume_text for skill in keywords):
            recommendations.append(career)

    if not recommendations:
        recommendations.append("general software engineering")

    return recommendations

import os
from PyPDF2 import PdfReader
from docx import Document
import re

def extract_text_from_resume(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text

    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text

    else:
        raise ValueError("Unsupported file format")

def get_recommendations(resume_text: str) -> list:
    resume_text = resume_text.lower()

    career_paths = {
        "data science": ["python", "machine learning", "data analysis", "pandas", "numpy", "scikit-learn", "tensorflow"],
        "web development": ["html", "css", "javascript", "react", "angular", "node", "frontend", "backend"],
        "android development": ["kotlin", "java", "android studio", "xml"],
        "cloud engineer": ["aws", "azure", "gcp", "cloud", "docker", "kubernetes", "devops"],
        "ai engineer": ["deep learning", "nlp", "transformers", "pytorch", "tensorflow", "huggingface", "llm"]
    }

    recommendations = []

    for career, keywords in career_paths.items():
        if any(skill in resume_text for skill in keywords):
            recommendations.append(career)

    if not recommendations:
        recommendations.append("general software engineering")

    return recommendations

